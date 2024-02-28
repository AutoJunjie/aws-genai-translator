import concurrent.futures
import docx
#from translator import agent_bedrock
import time
import logging
import os
import boto3
import json

TABLE_NAME = os.getenv('TABLE_NAME', '')
PRIMARY_KEY = os.getenv('PRIMARY_KEY', '')
UploadBucket = os.getenv('UploadBucket', '')

s3 = boto3.client('s3')

logger = logging.getLogger('translator')
logger.setLevel(logging.INFO)

bedrock = boto3.client(service_name='bedrock-runtime')

CONST_CLINT_SIDE_ERROR = "client_side_error"


def _get_complete_prompt(query, from_language='Chinese', to_language='English'):
    full_prompt = f"""
            You are a helpful and honest AI assistant, now I want you to 
            help in translation for the give text. 
            you will translate the {from_language} to its {to_language} version.  
            The following are the rules to follow during the translation.
            
            * The input will be in <TO_TRANSLATE> tag. they can be words, numbers, or single character, 
                Sometimes they are already in the target language, then only respond the original text into the <TRANSLATED> tag.
            * it is OK if you don't very confident to translate, 
                in such cases, you can give the best translate you can, because we will have human review later on.
            * Your output will be put to <TRANSLATED></TRANSLATED> tag. 
            * So, in summary, <TRANSLATED> tag should contain translated or original text, 
                <error> tag should contain the reason why you cannot translate.
            
            <TO_TRANSLATE>{query}</TO_TRANSLATE>
        """
    return json.dumps({
        "prompt": f"\n\nHuman:{full_prompt} \n\nAssistant:",
        "max_tokens_to_sample": 4096,
        "temperature": 0.5,
        "top_p": 1,
    })


def agent_bedrock(query, from_language, to_language):
    """
    Send the query to bedrock.
    """
    # modelId = 'anthropic.claude-v2:1'
    modelId = 'anthropic.claude-v2'
    accept = 'application/json'
    contentType = 'application/json'

    response = bedrock.invoke_model(body=_get_complete_prompt(
        query, from_language, to_language), modelId=modelId, accept=accept, contentType=contentType)
    response_body = json.loads(response.get('body').read())
    final_response = response_body['completion']

    logger.info(f"< Bedrock response: {final_response}")
    match = re.search(r'<TRANSLATED>([\s\S]*?)</TRANSLATED>', final_response)
    if match:
        final_response = match.group(1)
    logger.info(f"> input is: {query}, from {from_language} to {to_language} \n < Final response: {final_response}\n\n")
    return final_response

def agent_bedrock_streaming(query, from_language, to_language):
    """
    Send the query to bedrock.
    TODO: handle the bedrock error and rejections from bedrock.
    """
    # modelId = 'anthropic.claude-v2:1'
    modelId = 'anthropic.claude-v2'
    accept = 'application/json'
    contentType = 'application/json'
    
    logger.info(f'> input is: {query}, from {from_language} to {to_language}')

    response = bedrock.invoke_model_with_response_stream(body=_get_complete_prompt(
        query, from_language, to_language), modelId=modelId, accept=accept, contentType=contentType)
    # logger.info(response)
    response_body = response.get('body')

    # Iterate over events in the event stream as they come
    final_response = ""
    for event in response_body:
        # If we received a records event, write the data to a file
        # logger.info(f">> {event}")
        data = json.loads(event['chunk']['bytes'])
        # logger.info(data['completion'], end='')

        final_response = final_response + data['completion']
        # If we received a progress event, logger.info the details
        if data['stop_reason'] is not None:
            logger.info(">>>> stop reason: "+ data['stop_reason'])
            continue

    logger.info(f"< Bedrock response: {final_response}")
    match = re.search(r'<TRANSLATED>([\s\S]*?)</TRANSLATED>', final_response)
    if match:
        final_response = match.group(1)
    logger.info(f"< Final response: {final_response}\n\n")
    return final_response


def update_item(key, update_expression, expression_values, expression_keys):
    # Create a DynamoDB resource
    dynamodb = boto3.resource('dynamodb')

    # Get the table
    table = dynamodb.Table(TABLE_NAME)

    # Update the item
    response = table.update_item(
        Key=key,
        UpdateExpression=update_expression,
        ExpressionAttributeValues=expression_values,
        ExpressionAttributeNames=expression_keys
    )

    return response

def handler(event, context):
    logging.basicConfig(filename='app.log', filemode='w', level=logging.INFO,
                        format='%(name)s - %(levelname)s - %(message)s')

    logger = logging.getLogger('translator')
    logger.setLevel(logging.INFO)

    print(event)



    body = event.get('Records', [{}])[0].get('dynamodb', {}).get('NewImage', {})
    itemId = body.get('id', {}).get('S')
    eventName =  event.get('Records', [{}])[0].get('eventName')
    SourceKey = body.get('sourceKey', {}).get('S')

    if eventName != "INSERT":
        return {
            'statusCode': 200,
            'body': f"Not a INSERT event, nothing to do, session aborted'"
        }

    object_key = body.get('sourceKey', {}).get('S')
    #bucket_name = body.get('bucket_name', {}).get('S')
    FROM_LANGUAGE = body.get('languageSource', {}).get('S')
    print(body.get('languageTargets'))
    TO_LANGUAGE = body.get('languageTargets', {}).get('L')[0].get('S')
    #status = body.get('status', {}).get('S')
    print(f"********************{object_key}")
    s3_file_name = object_key.split("/")[-1]
    local_file_path = f'/tmp/{s3_file_name}'

    #FROM_LANGUAGE = 'Chinese'
    #TO_LANGUAGE = 'English'
    bucket_name = UploadBucket
    #object_key = 'translation/test6.docx'
    #s3_file_name = object_key.split("/")[-1]
    #local_file_path = f'/tmp/{s3_file_name}'

    # Replace 'bucket-name' and 'file-name.docx' with your bucket name and file name
    s3.download_file(bucket_name, object_key, local_file_path)

    doc = docx.Document(local_file_path)

    def translate(text):
        text = text.strip()
        if text:
            return agent_bedrock(text, FROM_LANGUAGE, to_language=TO_LANGUAGE)
        logger.info('-- skip')
        return ''

    texts = []
    for para in doc.paragraphs:
        if para.text:
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame'):
            text_frame = shape.text_frame
            for para in text_frame.paragraphs:
                texts.append(para.text)
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            texts.append(para.text)

    logger.info(f">>>> Number of tasks: {len(texts)}")
    #
    print(f"*****{itemId}")
    update_item(
         key={PRIMARY_KEY: itemId},
         update_expression="SET #jobStatus=:jobStatus",
         expression_values={':jobStatus': 'PROCESSING'},
         expression_keys={'#jobStatus': 'jobStatus'}
    )
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        results = executor.map(translate, texts)

    translated_texts = []
    for result in results:
        translated_texts.append(result)

    i = 0
    for para in doc.paragraphs:
        if para.text:
            logger.info(f"assemble: {para.text}\n {translated_texts[i]}")
            para.text = translated_texts[i] 
            i += 1

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = translated_texts[i]
                i += 1

    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame'): 
            text_frame = shape.text_frame
            for para in text_frame.paragraphs:
                para.text = translated_texts[i]
                i += 1

    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            para.text = translated_texts[i]
            i += 1

    ts = time.strftime('%Y%m%d%H%M%S')
    s3_file_name = object_key.split("/")[-1]
    doc.save(os.path.join('/tmp', f'{s3_file_name}'))
    #save doc to s3
    output_key = SourceKey.replace("upload","output").replace(".docx", f"_{ts}.docx")
    print(f"outputkey:{output_key}")
    s3.upload_file(f'/tmp/{s3_file_name}', bucket_name, output_key)

    logger.info(f"<<< {ts} Completed.")
    #update_item(
    #     key={PRIMARY_KEY: itemId},
    #     update_expression="SET #attribute_name1=:jobStatus",
    #     expression_values={':jobStatus': 'COMPLETED'},
    #     expression_keys={'#jobStatus': 'jobStatus'}
    #)

    update_item(
         key={PRIMARY_KEY: itemId},
         update_expression='SET #js = :js, #tk = :tk',
         expression_values={
            ':js': 'COMPLETED',
            ':tk': {'langen': output_key}  # This is a list containing one string
        },
         expression_keys={
            '#js': 'jobStatus',
            '#tk': 'translateKey'
        }
    )

    logging.shutdown()

    #return {
    #    'statusCode': 200,
    #    'body': f"Translation completed. Document saved as 'test_4_en_{ts}.docx'"
    #}
    
    return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'OPTIONS,POST'
            },
            'body': json.dumps(f"Received action")
            }